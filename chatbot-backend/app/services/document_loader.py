from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from app.utils.text_cleaner import clean_text

from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.config import settings


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def load_pdf(path: str) -> Optional[str]:
    try:
        doc = fitz.open(path)
        texts = []

        for page in doc:
            texts.append(page.get_text())

        doc.close()
        return clean_text("\n".join(texts))

    except Exception as e:
        print(f"[ERROR] Gagal membaca PDF {path}: {e}")
        return None


def load_txt(path: str) -> Optional[str]:
    try:
        with open(path, "r", encoding="utf-8") as file:
            return clean_text(file.read())

    except Exception as e:
        print(f"[ERROR] Gagal membaca TXT {path}: {e}")
        return None


def load_docx(path: str) -> Optional[str]:
    try:
        doc = Document(path)
        texts = []

        for paragraph in doc.paragraphs:
            texts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)

        return clean_text("\n".join(texts))

    except Exception as e:
        print(f"[ERROR] Gagal membaca DOCX {path}: {e}")
        return None


def load_document(path: str) -> Optional[str]:
    file_path = Path(path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return load_pdf(str(file_path))

    if ext == ".txt":
        return load_txt(str(file_path))

    if ext == ".docx":
        return load_docx(str(file_path))

    print(f"[WARNING] Format tidak didukung: {file_path.name}")
    return None


def scan_knowledge_base(folder_path: str) -> list[Path]:
    base_path = Path(folder_path)

    if not base_path.exists():
        print(f"[WARNING] Folder knowledge base tidak ditemukan: {folder_path}")
        return []

    files = [
        file
        for file in base_path.rglob("*")
        if file.is_file() and file.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    return files

def chunk_text(text: str, source_file: str) -> list[dict]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""],
    )

    chunks = splitter.split_text(text)

    chunks_with_metadata = []

    for index, chunk in enumerate(chunks):
        chunks_with_metadata.append(
            {
                "source_file": source_file,
                "chunk_index": index,
                "text": chunk,
            }
        )

    return chunks_with_metadata


def load_and_chunk_knowledge_base(folder_path: str) -> list[dict]:
    files = scan_knowledge_base(folder_path)
    all_chunks = []

    for file in files:
        text = load_document(str(file))

        if not text:
            print(f"[WARNING] File dilewati karena tidak memiliki teks: {file.name}")
            continue

        chunks = chunk_text(text, file.name)
        all_chunks.extend(chunks)

        print(f"[INFO] {file.name}: {len(chunks)} chunk")

    return all_chunks