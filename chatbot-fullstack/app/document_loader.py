from __future__ import annotations

from pathlib import Path
from typing import Optional

import fitz
from docx import Document

from app.config import settings
from app.text_cleaner import clean_text


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def load_pdf(path: str) -> Optional[str]:
    try:
        doc = fitz.open(path)
        texts = [page.get_text() for page in doc]
        doc.close()
        return clean_text("\n".join(texts))
    except Exception as exc:
        print(f"[ERROR] Gagal membaca PDF {path}: {exc}")
        return None


def load_txt(path: str) -> Optional[str]:
    try:
        with open(path, "r", encoding="utf-8") as file:
            return clean_text(file.read())
    except Exception as exc:
        print(f"[ERROR] Gagal membaca TXT {path}: {exc}")
        return None


def load_docx(path: str) -> Optional[str]:
    try:
        doc = Document(path)
        texts = [paragraph.text for paragraph in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)

        return clean_text("\n".join(texts))
    except Exception as exc:
        print(f"[ERROR] Gagal membaca DOCX {path}: {exc}")
        return None


def load_document(path: str) -> Optional[str]:
    file_path = Path(path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return load_pdf(str(file_path))
    if ext == ".txt":
        return load_txt(str(file_path))
    if ext == ".docx":
        return load_docx(str(file_path))

    print(f"[WARNING] Format tidak didukung: {file_path.name}")
    return None


def scan_knowledge_base(folder_path: str) -> list[Path]:
    base_path = Path(folder_path)

    if not base_path.exists():
        print(f"[WARNING] Folder knowledge base tidak ditemukan: {folder_path}")
        return []

    return [
        file
        for file in sorted(base_path.rglob("*"))
        if file.is_file() and file.suffix.lower() in SUPPORTED_EXTENSIONS
    ]


def chunk_text(text: str, source_file: str) -> list[dict]:
    size = settings.CHUNK_SIZE
    overlap = settings.CHUNK_OVERLAP

    if size <= 0:
        size = 1000
    if overlap < 0:
        overlap = 0
    if overlap >= size:
        overlap = max(0, size // 5)

    chunks_with_metadata: list[dict] = []
    start = 0
    index = 0

    while start < len(text):
        end = min(len(text), start + size)
        chunk = text[start:end].strip()

        if chunk:
            chunks_with_metadata.append(
                {
                    "source_file": source_file,
                    "chunk_index": index,
                    "text": chunk,
                }
            )
            index += 1

        if end >= len(text):
            break

        start = max(0, end - overlap)

    return chunks_with_metadata


def load_and_chunk_knowledge_base(folder_path: str) -> list[dict]:
    files = scan_knowledge_base(folder_path)
    all_chunks: list[dict] = []

    for file in files:
        text = load_document(str(file))
        if not text:
            print(f"[WARNING] File dilewati karena tidak memiliki teks: {file.name}")
            continue

        chunks = chunk_text(text, file.name)
        all_chunks.extend(chunks)
        print(f"[INFO] {file.name}: {len(chunks)} chunk")

    return all_chunks

